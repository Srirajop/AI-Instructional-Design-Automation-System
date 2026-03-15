import io
import re
import html
from datetime import datetime
from typing import Dict
from docx import Document
from docx.shared import RGBColor as DocxRGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

def export_design_doc_to_xlsx(design_doc_text: str, intake_data: Dict) -> io.BytesIO:
    """Export Design Document to Excel matching the template exactly."""
    try:
        output = io.BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.title = "Design Document"
        
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        yellow_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")
        grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        ws.merge_cells('A1:G1')
        cell = ws['A1']
        cell.value = f"Training Program For {intake_data.get('business_unit', 'Organization')}"
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF", size=14)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 25
        
        ws['A2'] = "Trainers"
        ws['A2'].font = Font(bold=True, size=10, color="0000FF")
        ws['A3'] = "Instructor / AI Assistant"
        ws['A3'].alignment = Alignment(wrap_text=True, vertical='top')
        ws.row_dimensions[3].height = 30
        
        ws.merge_cells('A4:G4')
        cell = ws['A4']
        cell.value = f"Design Document for {intake_data.get('course_title', 'Your Course')} — Self-paced eLearning"
        cell.fill = yellow_fill
        cell.font = Font(bold=True, size=11)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[4].height = 25
        
        headers = ["Module", "Delivery Mode", "Learning Objectives", "Topics", 
                   "Recommended Strategy", "Activities/Assessment", "Duration"]
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col_num)
            cell.value = header
            cell.fill = grey_fill
            cell.font = Font(bold=True, size=10)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border
        ws.row_dimensions[5].height = 30
        
        if not design_doc_text:
            wb.save(output) # Ensure valid empty file
            output.seek(0)
            return output
            
        row_num = 6
        lines = design_doc_text.split('\n')
        in_table = False
        
        table_row_count = 0
        for line in lines:
            line_s = line.strip()
            # 1. Detect Table Start
            # Looking for a line with at least 3 pipes OR identifying headers
            row_has_pipes = line_s.count('|') >= 2
            is_header_row = any(k.lower() in line_s.lower() for k in ['Module', 'Delivery', 'Objectives', 'Topics', 'Strategy'])
            
            if not in_table and row_has_pipes and is_header_row:
                in_table = True
                table_row_count = 0
                continue # Skip header row
            
            if in_table:
                # 2. Skip Separator or Empty Header Repeats
                if '---' in line_s and row_has_pipes:
                     continue
                
                # 3. Detect Data Row
                if row_has_pipes:
                    table_row_count += 1
                    
                    # Robust Pipe Splitting
                    row_content = line_s
                    if row_content.startswith('|'): row_content = row_content[1:]
                    if row_content.endswith('|'): row_content = row_content[:-1]
                    
                    cells = [c.strip() for c in row_content.split('|')]
                    
                    # Fill cells matching our 7-column template
                    if len(cells) > 0:
                        for col_idx in range(1, 8):
                            # Get content if exists, else empty string
                            content = cells[col_idx-1] if (col_idx-1) < len(cells) else ""
                            
                            clean_content = html.unescape(content)
                            clean_content = clean_content.replace('<br>', '\n').replace('•', '•').replace('**', '').replace('__', '')
                            # Preserve bullet points better
                            clean_content = clean_content.replace('&nbsp;', ' ')
                            clean_content = re.sub(r'<[^>]+>', '', clean_content)
                            
                            cell = ws.cell(row=row_num, column=col_idx)
                            cell.value = clean_content.strip()
                            cell.alignment = Alignment(vertical='top', wrap_text=True)
                            cell.border = thin_border
                        
                        ws.row_dimensions[row_num].height = 80 # Default row height
                        row_num += 1
                
                # 4. Detect Table End (Optional)
                elif table_row_count > 0 and line_s and not row_has_pipes:
                    in_table = False
        
        column_configs = [
            ('A', 35), ('B', 20), ('C', 50), ('D', 50),
            ('E', 50), ('F', 40), ('G', 15)
        ]
        for col_letter, width in column_configs:
            ws.column_dimensions[col_letter].width = width
            
        for row in ws.iter_rows(min_row=6, max_row=row_num-1):
            max_height = 60
            for cell in row:
                cell.alignment = Alignment(vertical='top', wrap_text=True)
                if cell.value:
                    col_letter = cell.column_letter
                    col_width = ws.column_dimensions[col_letter].width
                    effective_width = col_width * 0.9
                    if effective_width > 0:
                        val_str = str(cell.value)
                        lines_wrapped = len(val_str) / effective_width
                        lines_newlines = val_str.count('\n')
                        total_lines = lines_wrapped + lines_newlines + 1
                        cell_height = total_lines * 15
                        if cell_height > max_height:
                            max_height = cell_height
            ws.row_dimensions[row[0].row].height = max_height
        
        wb.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        print(f"Error exporting to Excel: {e}")
        return None

def export_storyboard_to_docx(storyboard_text: str, intake_data: Dict) -> io.BytesIO:
    """Export Storyboard to Word format with Table support."""
    try:
        output = io.BytesIO()
        doc = Document()
        
        title = doc.add_heading(f"eLearning Storyboard", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(intake_data.get('course_title', 'Untitled Course'), level=1)
        doc.add_paragraph(f"Business Unit: {intake_data.get('business_unit', '')}")
        doc.add_paragraph(f"Target Audience: {intake_data.get('target_audience', '')}")
        doc.add_paragraph(f"Interactivity: {intake_data.get('interactivity_level', '')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph("")
        
        if not storyboard_text:
            doc.save(output) # Ensure valid empty file
            output.seek(0)
            return output
            
        lines = storyboard_text.split('\n')
        table_lines = []
        in_table = False
        headers = []
        seen_headers = set()

        for line in lines:
            line_stripped = line.strip()
            
            # 1. Detect Table Start or Row
            row_has_pipes = line_stripped.count('|') >= 2
            is_header_row = any(k.lower() in line_stripped.lower() for k in ["Section", "Topics", "Visual", "On-screen", "Audio", "On-Screen", "OST"])
            
            if row_has_pipes:
                # 2. Skip Separator Lines
                if re.search(r'^[|:\s-]*$', line_stripped) or '| ---' in line_stripped or '|:---' in line_stripped:
                    continue
                
                if not in_table:
                    in_table = True
                    # Clean up pipes and split
                    header_line = line_stripped
                    if header_line.startswith('|'): header_line = header_line[1:]
                    if header_line.endswith('|'): header_line = header_line[:-1]
                    headers = [h.strip() for h in header_line.split('|')]
                    table_lines = []
                else:
                    cells_line = line_stripped
                    if cells_line.startswith('|'): cells_line = cells_line[1:]
                    if cells_line.endswith('|'): cells_line = cells_line[:-1]
                    cells = [c.strip() for c in cells_line.split('|')]
                    table_lines.append(cells)
                continue
            
            if in_table and not line_stripped.startswith('|'):
                if headers and table_lines:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        if i < len(hdr_cells):
                            hdr_cells[i].text = header
                            hdr_cells[i].paragraphs[0].runs[0].bold = True
                    for row_data in table_lines:
                        row = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            if i < len(row):
                                clean_text = html.unescape(cell_text).replace('<br>', '\n').replace('**', '')
                                row[i].text = clean_text
                in_table = False
                headers = []
                table_lines = []

            if line_stripped.startswith('===') or line_stripped.startswith('---'):
                continue
                
            clean_line = line_stripped.replace('**', '').strip()
            
            if clean_line.upper().startswith('MODULE'):
                if clean_line.upper() in seen_headers:
                    continue
                seen_headers.add(clean_line.upper())
                p = doc.add_heading('', level=1)
                run = p.add_run(clean_line)
                run.bold = True
                run.font.color.rgb = DocxRGBColor(31, 78, 120)
            elif clean_line.upper().startswith('SCREEN'):
                if clean_line.upper() in seen_headers:
                    continue
                seen_headers.add(clean_line.upper())
                p = doc.add_heading('', level=2)
                run = p.add_run(clean_line)
                run.bold = True
                run.font.color.rgb = DocxRGBColor(0, 0, 0)
            elif ':' in clean_line:
                parts = clean_line.split(':', 1)
                key = parts[0].strip()
                val = parts[1].strip() if len(parts) > 1 else ""
                
                is_header_key = (key.isupper() and len(key) < 40) or \
                                any(k in key for k in ["Topic", "Date", "Target", "Business", "Interactivity"])
                
                if is_header_key:
                    p = doc.add_paragraph()
                    run = p.add_run(key + ':')
                    run.bold = True
                    run.font.color.rgb = DocxRGBColor(31, 78, 120)
                    if val:
                        p.add_run(' ' + val)
                else:
                    doc.add_paragraph(clean_line)
            elif clean_line:
                doc.add_paragraph(clean_line)
        
        if in_table and headers and table_lines:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = header
            for row_data in table_lines:
                row = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    if i < len(row):
                        row[i].text = html.unescape(cell_text).replace('<br>', '\n').replace('**', '')

        doc.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        print(f"Error exporting to Docx: {e}")
        output = io.BytesIO()
        doc = Document()
        doc.add_heading("Storyboard Export (Fallback)", 0)
        doc.add_paragraph(storyboard_text)
        doc.save(output)
        output.seek(0)
        return output
